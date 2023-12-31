import os
import shutil
import subprocess
import pandas as pd
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import glob
import sys
import configparser

# Check if a command line argument has been provided
if len(sys.argv) < 2:
    print("Usage: python _analyze_general.py <config_file>")
    sys.exit(1)

# The second command line argument is expected to be the config file name
config_file = sys.argv[1]

# Initialize the configparser
config = configparser.ConfigParser()

# Read the configuration file passed as a command line argument
config.read(config_file)

MAX = config.getint('DEFAULT', 'PLAN_SEARCH_TIMEOUT_SECONDS')

# List of scripts to copy
scripts = [
    '_analyze.py',
    '_generate_summary.py',
    '_normalize.py',
    '_pressure.py',
    '_report_root_folder_full.py',
    '_xyplot.py'
]

# Function to copy scripts to target folder and run _analyze.py
def process_folder(folder_name):
    for script in scripts:
        # Copy each script to the target folder
        shutil.copy(os.path.join(root_dir, script), os.path.join(folder_name, script))

    # Change to the target folder to run _analyze.py
    os.chdir(folder_name)
    subprocess.run(['python3', '_analyze.py', str(MAX)])
    os.chdir(root_dir)  # Change back to root directory

# Main process
root_dir = os.getcwd()  # Assuming scripts are located in the current working directory

# Automatically get all folders starting with "archive_pfile"
archive_folders = glob.glob('archive_pfile_*')  # This will return a list of all folders starting with "archive_pfile"

# Process each archive folder
for folder_name in archive_folders:
    process_folder(folder_name)

# Create summary folder and process each archive folder again to move generated files
summary_dir = os.path.join(root_dir, 'summary')
os.makedirs(summary_dir, exist_ok=True)

# Moving generated files to a summary directory
for folder_name in archive_folders:
    # Here, you don't create a new folder name based on the index, but rather use the existing folder name
    new_folder_name = folder_name.replace('archive_', '')  # Optional: remove 'archive_' prefix for summary folder names
    new_folder_path = os.path.join(summary_dir, new_folder_name)
    os.makedirs(new_folder_path, exist_ok=True)
    try:
        # Print the current working directory for debugging
        print("Current Working Directory: ", os.getcwd())

        # Copy generated files to the new folders in the summary directory
        for file_name in ['disp_vs_nodisp_plot.png', 'summary.csv']:
            src_path = os.path.join(folder_name, file_name)

            # Check if the source file exists
            if os.path.exists(src_path):
                shutil.copy(src_path, os.path.join(new_folder_path, file_name))
            else:
                print(f"File not found: {src_path}")
    except Exception as e:
        print("An error occurred:", e)

def generate_plots(summary_csv_path):
    # Load the data
    data = pd.read_csv(summary_csv_path)

    # Data preparation
    labels = data['Unnamed: 0']
    avg_time_disp = data['Average solution time for disp on permutations solved by both disp and nodisp']
    avg_time_nodisp = data['Average solution time for nodisp on permutations solved by both disp and nodisp']
    both_solved = data['Number of permutations solved by both disp and nodisp']
    only_disp_solved = data['Number of permutations solved only by disp']
    only_nodisp_solved = data['Number of permutations solved only by nodisp']

    # --- Visualization 1: Average Solution Time Comparison ---
    plt.figure(figsize=(14, 7))
    width = 0.35
    x = range(len(labels))
    x_disp = [i - width/2 for i in x]
    x_nodisp = [i + width/2 for i in x]
    plt.bar(x_disp, avg_time_disp, width, label='Disp', color='green')
    plt.bar(x_nodisp, avg_time_nodisp, width, label='NoDisp', color='orange')
    plt.xlabel('Problem Files')
    plt.ylabel('Average Solution Time (in units)')
    plt.title('Average Solution Time Comparison Between Disp and NoDisp')
    plt.xticks(ticks=x, labels=labels, rotation=45)
    plt.legend()
    plt.tight_layout()
    plt.savefig('plot1.png')  # Save the plot to a file

    # --- Visualization 2: Comparison of Permutations Solved by Method ---
    plt.figure(figsize=(12, 6))
    fig, ax = plt.subplots()
    rects1 = ax.bar(x, both_solved, width, label='Both', color='#1f77b4')
    rects2 = ax.bar(x, only_disp_solved, width, label='Only Disp', bottom=both_solved, color='#2ca02c')
    rects3 = ax.bar(x, only_nodisp_solved, width, label='Only NoDisp', bottom=both_solved+only_disp_solved, color='#ff7f0e')
    ax.set_xlabel('Problem Files')
    ax.set_ylabel('Number of Permutations Solved')
    ax.set_title('Comparison of Permutations Solved by Method')
    ax.set_xticks(x)
    ax.set_xticklabels(labels, rotation=45)
    ax.legend()
    plt.tight_layout()
    plt.savefig('plot2.png')  # Save the plot to a file

def generate_doc(summary_dir):
    doc = Document()
    doc.add_heading('Analysis Report', 0)

    # Add global plots to the document
    for plot_name in ['plot1.png', 'plot2.png']:
        plot_path = os.path.join(summary_dir, plot_name)
        if os.path.exists(plot_path):
            doc.add_picture(plot_path, width=Inches(6))
        else:
            print(f"Global plot not found: {plot_path}")

    for folder_name in archive_folders:
        new_folder_name = folder_name.replace('archive_', '')
        new_folder_path = os.path.join(summary_dir, new_folder_name)

        # Add CSV data to document
        csv_file = os.path.join(new_folder_path, 'summary.csv')
        if os.path.exists(csv_file):
            df = pd.read_csv(csv_file)
            table = doc.add_table(df.shape[0] + 1, df.shape[1])
            for j, (index, row) in enumerate(df.iterrows()):
                for k, value in enumerate(row):
                    table.cell(j + 1, k).text = str(value)
            for k, name in enumerate(df.columns):
                table.cell(0, k).text = name
        else:
            print(f"CSV file not found: {csv_file}")

        # Add plot to document
        plot_file = os.path.join(new_folder_path, 'disp_vs_nodisp_plot.png')
        if os.path.exists(plot_file):
            doc.add_picture(plot_file, width=Inches(6))
        else:
            print(f"Folder plot not found: {plot_file}")

    doc_path = os.path.join(summary_dir, 'summary.docx')
    doc.save(doc_path)

    return doc_path

def generate_summary_csv(summary_dir):
    all_data = []  # List to hold data from all CSV files

    for folder_name in archive_folders:
        new_folder_name = folder_name.replace('archive_', '')  # Same modification as above
        new_folder_path = os.path.join(summary_dir, new_folder_name)

        # Read CSV data
        csv_file = os.path.join(new_folder_path, 'summary.csv')
        df = pd.read_csv(csv_file, index_col=0)  # Assuming the descriptions are in the first column
        df_transposed = df.T  # Transpose the DataFrame
        df_transposed.index = [new_folder_name]  # Set the index to the pfile name
        all_data.append(df_transposed)

    # Concatenate all data into one DataFrame along rows
    summary_df = pd.concat(all_data)

    # Write summary data to a new CSV file
    summary_csv_path = os.path.join(summary_dir, 'summary_table.csv')
    summary_df.to_csv(summary_csv_path)

    return summary_csv_path

# After processing all folders or at the end of your script
summary_csv_path = generate_summary_csv(summary_dir)

# Generate plots
generate_plots(summary_csv_path)

# After copying files to the new folders in the summary directory
doc_path = generate_doc(summary_dir)

# Open the document
os.system(f'xdg-open {doc_path}')

# Optionally, open the summary CSV file
subprocess.run(['xdg-open', summary_csv_path])
