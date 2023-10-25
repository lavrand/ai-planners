import os
import shutil
import subprocess
import pandas as pd
from docx import Document
from docx.shared import Inches
import subprocess

# Get the total number of archive folders (assuming they are sequentially named and start from 1)
N = 22  # Set this to the actual number of archive folders you have

# List of scripts to copy
scripts = [
    '_analyze.py',
    '_generate_summary.py',
    '_normalize.py',
    '_pressure.py',
    '_report_root_folder_full.py',
    '_xyplot.py'
]

# Function to copy scripts to target folder and run _analyze.py
def process_folder(folder_name):
    for script in scripts:
        # Copy each script to the target folder
        shutil.copy(os.path.join(root_dir, script), os.path.join(folder_name, script))

    # Change to the target folder to run _analyze.py
    os.chdir(folder_name)
    subprocess.run(['python3', '_analyze.py'])
    os.chdir(root_dir)  # Change back to root directory


# Main process
root_dir = os.getcwd()  # Assuming scripts are located in the current working directory

# Process each archive folder
for i in range(1, N + 1):
    folder_name = f'archive_pfile_{i}'
    process_folder(folder_name)

# Create summary folder and process each archive folder again to move generated files
summary_dir = os.path.join(root_dir, 'summary')
os.makedirs(summary_dir, exist_ok=True)

for i in range(1, N + 1):
    folder_name = f'archive_pfile_{i}'
    new_folder_name = f'pfile_{i}'
    new_folder_path = os.path.join(summary_dir, new_folder_name)
    os.makedirs(new_folder_path, exist_ok=True)

    # Copy generated files to the new folders in the summary directory
    for file_name in ['disp_vs_nodisp_plot.png', 'summary.csv']:
        shutil.copy(os.path.join(folder_name, file_name), os.path.join(new_folder_path, file_name))


def generate_doc(summary_dir):
    doc = Document()

    for i in range(1, N + 1):
        new_folder_name = f'pfile_{i}'
        new_folder_path = os.path.join(summary_dir, new_folder_name)

        # Add CSV data to document
        csv_file = os.path.join(new_folder_path, 'summary.csv')
        df = pd.read_csv(csv_file)
        table = doc.add_table(df.shape[0] + 1, df.shape[1])
        for j, (index, row) in enumerate(df.iterrows()):
            for k, value in enumerate(row):
                table.cell(j + 1, k).text = str(value)
        for k, name in enumerate(df.columns):
            table.cell(0, k).text = name

        # Add plot to document
        plot_file = os.path.join(new_folder_path, 'disp_vs_nodisp_plot.png')
        doc.add_picture(plot_file, width=Inches(6))

    doc_path = os.path.join(summary_dir, 'summary.docx')
    doc.save(doc_path)

    return doc_path

# After copying files to the new folders in the summary directory
doc_path = generate_doc(summary_dir)

# Open the document
os.system(f'xdg-open {doc_path}')

def generate_summary_csv(summary_dir):
    all_data = []  # List to hold data from all CSV files

    for i in range(1, N + 1):
        new_folder_name = f'pfile_{i}'
        new_folder_path = os.path.join(summary_dir, new_folder_name)

        # Read CSV data
        csv_file = os.path.join(new_folder_path, 'summary.csv')
        df = pd.read_csv(csv_file, index_col=0)  # Assuming the descriptions are in the first column
        df_transposed = df.T  # Transpose the DataFrame
        df_transposed.index = [new_folder_name]  # Set the index to the pfile name
        all_data.append(df_transposed)

    # Concatenate all data into one DataFrame along rows
    summary_df = pd.concat(all_data)

    # Write summary data to a new CSV file
    summary_csv_path = os.path.join(summary_dir, 'summary_table.csv')
    summary_df.to_csv(summary_csv_path)

    return summary_csv_path

# After processing all folders or at the end of your script
summary_csv_path = generate_summary_csv(summary_dir)

# Optionally, open the summary CSV file
subprocess.run(['xdg-open', summary_csv_path])
