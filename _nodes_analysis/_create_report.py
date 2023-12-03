#!/usr/bin/python3
import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def add_table_from_csv(doc, file_path):
    df = pd.read_csv(file_path)

    # Apply rounding to numeric rows
    for col in df.columns:
        if df[col].dtype.kind in 'fi':  # Check for float or integer columns
            df[col] = df[col].round(0).astype(int)

    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'

    # Add the header rows.
    for j in range(len(df.columns)):
        table.cell(0, j).text = df.columns[j]

    # Add the rest of the data frame
    for i in range(df.shape[0]):
        row_cells = table.add_row().cells
        for j in range(df.shape[1]):
            value = df.iloc[i, j]
            row_cells[j].text = str(value)
            # Check for disp and nodisp columns and apply color
            if 'disp' in df.columns and 'nodisp' in df.columns:
                disp_index = df.columns.get_loc('disp')
                nodisp_index = df.columns.get_loc('nodisp')
                if j in [disp_index, nodisp_index] and pd.to_numeric(value, errors='coerce') is not None:
                    if df.iloc[i, disp_index] > df.iloc[i, nodisp_index]:
                        # Color red
                        row_cells[j]._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="FF0000"/>'.format(nsdecls('w'))))
                    elif df.iloc[i, disp_index] < df.iloc[i, nodisp_index]:
                        # Color green
                        row_cells[j]._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="00FF00"/>'.format(nsdecls('w'))))

def create_report(image_path, summary_csv, simulated_time_csv, report_docx):
    doc = Document()
    doc.add_picture(image_path, width=Inches(6))
    doc.add_page_break()
    add_table_from_csv(doc, summary_csv)
    doc.add_page_break()
    add_table_from_csv(doc, simulated_time_csv)
    doc.save(report_docx)

if __name__ == "__main__":
    create_report('disp_vs_nodisp_plot.png', 'summary.csv', 'Simulated_Time_Solution_Found.csv', 'report.docx')
